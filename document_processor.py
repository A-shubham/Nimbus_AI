# --- document_processor.py (Source Highlighting Version) ---

import os
import logging
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_vertexai import VertexAIEmbeddings
from langchain_google_vertexai.vectorstores import VectorSearchVectorStore
from google.cloud import aiplatform
# We need the Document class to add metadata
from langchain_core.documents import Document 

import config
from pypdf import PdfReader
from docx import Document as DocxDocument

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)

def _extract_text_from_file(file_path: str):
    """Extracts text and returns it with the filename as metadata."""
    text = ""
    filename = os.path.basename(file_path)
    ext = os.path.splitext(filename)[1].lower()

    try:
        if ext == ".pdf":
            reader = PdfReader(file_path)
            text = "".join(page.extract_text() + "\n" for page in reader.pages)
        elif ext == ".docx":
            doc = DocxDocument(file_path)
            text = "".join(para.text + "\n" for para in doc.paragraphs)
        elif ext == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
    except Exception as e:
        logging.error(f"Error extracting text from {filename}: {e}")
        return None
    
    # Return a Document object with content and metadata
    return Document(page_content=text, metadata={"source": filename})

def process_and_upload_documents(file_paths: list[str]):
    logging.info(f"--- Starting document processing for {len(file_paths)} files ---")
    
    # Extract text from all files into Document objects
    docs_with_metadata = [_extract_text_from_file(fp) for fp in file_paths]
    docs_with_metadata = [doc for doc in docs_with_metadata if doc is not None] # Filter out failed extractions
    
    if not docs_with_metadata:
        logging.warning("No text could be extracted. Aborting upload.")
        return False
    
    logging.info("Chunking documents...")
    # The splitter will preserve the metadata for each chunk
    chunks = text_splitter.split_documents(docs_with_metadata)
    logging.info(f"Generated {len(chunks)} text chunks with metadata.")

    if not chunks:
        logging.warning("Text extraction resulted in zero chunks. Aborting upload.")
        return False
    
    try:
        logging.info("Initializing Vertex AI Vector Search for upload...")
        vector_store = VectorSearchVectorStore.from_components(
            project_id=config.PROJECT_ID,
            region=config.REGION,
            gcs_bucket_name=config.GCS_BUCKET_NAME,
            index_name=config.VECTOR_SEARCH_INDEX_NAME,
            index_id=config.VECTOR_SEARCH_INDEX_ID,
            endpoint_id=config.VECTOR_SEARCH_ENDPOINT_ID,
            embedding=VertexAIEmbeddings(model_name="text-embedding-004"),
        )
        
        logging.info(f"Adding {len(chunks)} documents to Vertex AI index...")
        # Use add_documents to include metadata
        vector_store.add_documents(documents=chunks)
        logging.info("--- Successfully uploaded documents to Vertex AI Vector Search ---")
        return True
    except Exception as e:
        logging.error(f"Failed to upload documents to Vertex AI: {e}", exc_info=True)
        return False